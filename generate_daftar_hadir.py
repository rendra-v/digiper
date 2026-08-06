import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 90 Working Days Master List matching Logbook
DATES_90 = [
    ("Selasa", "03/02/2026", "WFO (Kantor MA)"),
    ("Rabu", "04/02/2026", "WFO (Kantor MA)"),
    ("Kamis", "05/02/2026", "WFO (Kantor MA)"),
    ("Jumat", "06/02/2026", "WFH (Work From Home)"),
    ("Sabtu", "07/02/2026", "WFH (Work From Home)"),
    ("Senin", "09/02/2026", "WFO (Kantor MA)"),
    ("Selasa", "10/02/2026", "WFO (Kantor MA)"),
    ("Rabu", "11/02/2026", "WFO (Kantor MA)"),
    ("Kamis", "12/02/2026", "WFO (Kantor MA)"),
    ("Jumat", "13/02/2026", "WFO (Kantor MA)"),
    ("Sabtu", "14/02/2026", "WFH (Work From Home)"),
    ("Senin", "16/02/2026", "WFO (Kantor MA)"),
    ("Selasa", "17/02/2026", "WFO (Kantor MA)"),
    ("Rabu", "18/02/2026", "WFO (Kantor MA)"),
    ("Kamis", "19/02/2026", "WFO (Kantor MA)"),
    ("Jumat", "20/02/2026", "WFH (Work From Home)"),
    ("Sabtu", "21/02/2026", "WFH (Work From Home)"),
    ("Senin", "23/02/2026", "WFO (Kantor MA)"),
    ("Selasa", "24/02/2026", "WFO (Kantor MA)"),
    ("Rabu", "25/02/2026", "WFO (Kantor MA)"),
    ("Kamis", "26/02/2026", "WFO (Kantor MA)"),
    ("Jumat", "27/02/2026", "WFO (Kantor MA)"),
    ("Sabtu", "28/02/2026", "WFH (Work From Home)"),
    ("Senin", "02/03/2026", "WFO (Kantor MA)"),
    ("Selasa", "03/03/2026", "WFO (Kantor MA)"),
    ("Rabu", "04/03/2026", "WFO (Kantor MA)"),
    ("Kamis", "05/03/2026", "WFO (Kantor MA)"),
    ("Jumat", "06/03/2026", "WFH (Work From Home)"),
    ("Sabtu", "07/03/2026", "WFH (Work From Home)"),
    ("Senin", "09/03/2026", "WFO (Kantor MA)"),
    ("Selasa", "10/03/2026", "WFO (Kantor MA)"),
    ("Rabu", "11/03/2026", "WFO (Kantor MA)"),
    ("Kamis", "12/03/2026", "WFO (Kantor MA)"),
    ("Jumat", "13/03/2026", "WFO (Kantor MA)"),
    ("Sabtu", "14/03/2026", "WFH (Work From Home)"),
    ("Senin", "16/03/2026", "WFO (Kantor MA)"),
    ("Selasa", "17/03/2026", "WFO (Kantor MA)"),
    ("Rabu", "18/03/2026", "WFO (Kantor MA)"),
    ("Kamis", "19/03/2026", "WFO (Kantor MA)"),
    ("Jumat", "20/03/2026", "WFH (Work From Home)"),
    ("Sabtu", "21/03/2026", "WFH (Work From Home)"),
    ("Senin", "23/03/2026", "WFO (Kantor MA)"),
    ("Selasa", "24/03/2026", "WFO (Kantor MA)"),
    ("Rabu", "25/03/2026", "WFO (Kantor MA)"),
    ("Kamis", "26/03/2026", "WFO (Kantor MA)"),
    ("Jumat", "27/03/2026", "WFO (Kantor MA)"),
    ("Sabtu", "28/03/2026", "WFH (Work From Home)"),
    ("Senin", "11/05/2026", "WFO (Kantor MA)"),
    ("Selasa", "12/05/2026", "WFO (Kantor MA)"),
    ("Rabu", "13/05/2026", "WFO (Kantor MA)"),
    ("Kamis", "14/05/2026", "WFO (Kantor MA)"),
    ("Jumat", "15/05/2026", "WFH (Work From Home)"),
    ("Sabtu", "16/05/2026", "WFH (Work From Home)"),
    ("Senin", "18/05/2026", "WFO (Kantor MA)"),
    ("Selasa", "19/05/2026", "WFO (Kantor MA)"),
    ("Rabu", "20/05/2026", "WFO (Kantor MA)"),
    ("Kamis", "21/05/2026", "WFO (Kantor MA)"),
    ("Jumat", "22/05/2026", "WFO (Kantor MA)"),
    ("Sabtu", "23/05/2026", "WFH (Work From Home)"),
    ("Senin", "01/06/2026", "WFO (Kantor MA)"),
    ("Selasa", "02/06/2026", "WFO (Kantor MA)"),
    ("Rabu", "03/06/2026", "WFO (Kantor MA)"),
    ("Kamis", "04/06/2026", "WFO (Kantor MA)"),
    ("Jumat", "05/06/2026", "WFH (Work From Home)"),
    ("Sabtu", "06/06/2026", "WFH (Work From Home)"),
    ("Senin", "22/06/2026", "WFO (Kantor MA)"),
    ("Selasa", "23/06/2026", "WFO (Kantor MA)"),
    ("Rabu", "24/06/2026", "WFO (Kantor MA)"),
    ("Kamis", "25/06/2026", "WFO (Kantor MA)"),
    ("Jumat", "26/06/2026", "WFO (Kantor MA)"),
    ("Sabtu", "27/06/2026", "WFH (Work From Home)"),
    ("Senin", "29/06/2026", "WFO (Kantor MA)"),
    ("Selasa", "30/06/2026", "WFO (Kantor MA)"),
    ("Rabu", "01/07/2026", "WFO (Kantor MA)"),
    ("Kamis", "02/07/2026", "WFO (Kantor MA)"),
    ("Jumat", "03/07/2026", "WFH (Work From Home)"),
    ("Sabtu", "04/07/2026", "WFH (Work From Home)"),
    ("Senin", "20/07/2026", "WFO (Kantor MA)"),
    ("Selasa", "21/07/2026", "WFO (Kantor MA)"),
    ("Rabu", "22/07/2026", "WFO (Kantor MA)"),
    ("Kamis", "23/07/2026", "WFO (Kantor MA)"),
    ("Jumat", "24/07/2026", "WFO (Kantor MA)"),
    ("Sabtu", "25/07/2026", "WFH (Work From Home)"),
    ("Senin", "27/07/2026", "WFO (Kantor MA)"),
    ("Selasa", "28/07/2026", "WFO (Kantor MA)"),
    ("Rabu", "29/07/2026", "WFO (Kantor MA)"),
    ("Kamis", "30/07/2026", "WFO (Kantor MA)"),
    ("Jumat", "31/07/2026", "WFH (Work From Home)"),
    ("Sabtu", "01/08/2026", "WFH (Work From Home)"),
    ("Senin", "03/08/2026", "WFO (Kantor MA)")
]

STUDENTS = [
    ("BE1", "Rifa Reza Fahlevi", "1402023060", "DAFTAR_HADIR_BE1_RIFA_REZA_FAHLEVI"),
    ("BE2", "Anggota Backend 2", "1402023061", "DAFTAR_HADIR_BE2"),
    ("FE1", "Anggota Frontend 1", "1402023062", "DAFTAR_HADIR_FE1"),
    ("FE2", "Anggota Frontend 2", "1402023063", "DAFTAR_HADIR_FE2")
]

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            borders.append(border)
        tblPr[0].append(borders)

def generate_md(role, name, npm, filename):
    md_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{filename}.md"
    content = f"# DAFTAR HADIR MAGANG MAHASISWA\n"
    content += f"**Tahun Ajaran:** 2025/2026  \n"
    content += f"**Instansi:** Mahkamah Agung Republik Indonesia  \n"
    content += f"**Nama Mahasiswa:** {name}  \n"
    content += f"**NPM:** {npm}  \n"
    content += f"**Peran:** {role}  \n"
    content += f"**Periode:** 03 Februari 2026 – 03 Agustus 2026 (90 Hari Hadir / 100% Hadir)  \n\n"
    content += "---\n\n"
    content += "| NO | NAMA MAHASISWA | NPM | TANGGAL HADIR | H | I | S | A | KETERANGAN |\n"
    content += "|:--:|:---|:---:|:---:|:--:|:--:|:--:|:--:|:---|\n"

    for idx, (day, date, ket) in enumerate(DATES_90, 1):
        content += f"| {idx} | {name} | {npm} | {day}, {date} | 1 | 0 | 0 | 0 | {ket} |\n"

    content += "\n---\n"
    content += "**Keterangan:**\n"
    content += "- **H**: Hadir (100% / 90 Hari)\n"
    content += "- **I**: Izin (0)\n"
    content += "- **S**: Sakit (0)\n"
    content += "- **A**: Alpa (0)\n\n"
    content += "**Catatan Kehadiran:**\n"
    content += "- **WFO (Work From Office)**: Hari Senin s.d. Kamis & Jumat Terjadwal (Kantor Mahkamah Agung RI)\n"
    content += "- **WFH (Work From Home)**: Hari Sabtu & Jumat Terjadwal\n\n"
    content += "Jakarta, 03 Agustus 2026  \n"
    content += "**Instruktur / Pembimbing Magang Instansi:**  \n\n\n"
    content += "__________________________________________  \n"
    content += "NIP. …………………………………………  \n"

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"Generated MD: {md_path}")

def generate_html(role, name, npm, filename):
    html_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{filename}.html"
    html = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>DAFTAR HADIR MAGANG - {name}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 30px; background-color: #f8f9fa; }}
        .container {{ max-width: 1100px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #1f497d; margin-bottom: 5px; font-size: 22px; }}
        h2 {{ text-align: center; color: #555; margin-top: 0; font-size: 15px; font-weight: normal; margin-bottom: 25px; }}
        .info {{ margin-bottom: 20px; font-size: 14px; line-height: 1.6; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 13px; }}
        th {{ background-color: #1f497d; color: white; padding: 8px; text-align: center; border: 1px solid #1f497d; }}
        td {{ padding: 7px 8px; border: 1px solid #dce4ec; text-align: center; }}
        td.left {{ text-align: left; }}
        tr:nth-child(even) {{ background-color: #f2f5f9; }}
        .wfh {{ color: #d97706; font-weight: bold; }}
        .wfo {{ color: #15803d; font-weight: bold; }}
        .footer {{ margin-top: 30px; float: right; text-align: center; width: 250px; font-size: 14px; }}
        .clear {{ clear: both; }}
    </style>
</head>
<body>
<div class="container">
    <h1>DAFTAR HADIR MAGANG MAHASISWA</h1>
    <h2>Mahkamah Agung Republik Indonesia — Tahun Ajaran 2025/2026</h2>
    <div class="info">
        <b>Nama Mahasiswa:</b> {name}<br>
        <b>NPM:</b> {npm}<br>
        <b>Peran:</b> {role}<br>
        <b>Total Kehadiran:</b> 90 Hari (100% Hadir)
    </div>
    <table>
        <thead>
            <tr>
                <th width="5%">NO</th>
                <th width="25%">NAMA MAHASISWA</th>
                <th width="15%">NPM</th>
                <th width="22%">TANGGAL HADIR</th>
                <th width="5%">H</th>
                <th width="5%">I</th>
                <th width="5%">S</th>
                <th width="5%">A</th>
                <th width="18%">KETERANGAN</th>
            </tr>
        </thead>
        <tbody>
"""
    for idx, (day, date, ket) in enumerate(DATES_90, 1):
        cls = "wfh" if "WFH" in ket else "wfo"
        html += f"""            <tr>
                <td>{idx}</td>
                <td class="left">{name}</td>
                <td>{npm}</td>
                <td>{day}, {date}</td>
                <td><b>1</b></td>
                <td>0</td>
                <td>0</td>
                <td>0</td>
                <td class="{cls}">{ket}</td>
            </tr>\n"""

    html += f"""        </tbody>
    </table>
    <div style="margin-top: 20px; font-size: 13px;">
        <b>Keterangan:</b> H = Hadir (90 Hari), I = Izin (0), S = Sakit (0), A = Alpa (0)<br>
        <b>Status Kerja:</b> WFO = Work From Office (Kantor MA) | WFH = Work From Home (Sabtu & Jumat Terjadwal)
    </div>
    <div class="footer">
        Jakarta, 03 Agustus 2026<br>
        <b>Instruktur Magang Instansi</b><br><br><br><br>
        ___________________________<br>
        NIP. …………………………………
    </div>
    <div class="clear"></div>
</div>
</body>
</html>"""

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"Generated HTML: {html_path}")

def generate_docx(role, name, npm, filename):
    docx_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{filename}.docx"
    doc = docx.Document()
    
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR HADIR MAGANG MAHASISWA\n")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    r2 = p.add_run("Mahkamah Agung Republik Indonesia — Tahun Ajaran 2025/2026\n")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Meta
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(8)
    p_meta.add_run(f"Nama Mahasiswa : {name}\nNPM             : {npm}\nPeran           : {role}\nTotal Kehadiran : 90 Hari (100% Hadir)")

    # Table
    table = doc.add_table(rows=len(DATES_90) + 1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B0C4DE")

    headers = ["NO", "NAMA MAHASISWA", "NPM", "TANGGAL HADIR", "H", "I", "S", "A", "KETERANGAN"]
    col_widths = [Inches(0.4), Inches(1.8), Inches(1.0), Inches(1.4), Inches(0.3), Inches(0.3), Inches(0.3), Inches(0.3), Inches(1.6)]

    # Header Row
    for c_idx, h_text in enumerate(headers):
        cell = table.cell(0, c_idx)
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        cell.width = col_widths[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data Rows
    for r_idx, (day, date, ket) in enumerate(DATES_90, 1):
        row_vals = [str(r_idx), name, npm, f"{day}, {date}", "1", "0", "0", "0", ket]
        bg_color = "F2F5F9" if r_idx % 2 == 1 else "FFFFFF"

        for c_idx, val in enumerate(row_vals):
            cell = table.cell(r_idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]

            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(val)
            run.font.size = Pt(8.5)

            if "WFH" in val:
                run.bold = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Orange WFH
            elif c_idx == 4:
                run.bold = True
                run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D) # Green Hadir

    # Legend & Signature
    doc.add_paragraph()
    p_leg = doc.add_paragraph()
    p_leg.add_run("Keterangan: H = Hadir (90 Hari), I = Izin (0), S = Sakit (0), A = Alpa (0)\nStatus Kerja: WFO = Work From Office (Kantor MA) | WFH = Work From Home (Sabtu & Jumat Terjadwal)")
    p_leg.runs[0].font.size = Pt(9)

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.paragraph_format.space_before = Pt(20)
    p_sig.add_run("Jakarta, 03 Agustus 2026\nInstruktur Magang Instansi\n\n\n\n_________________________________\nNIP. ……………………………………………")
    p_sig.runs[0].font.size = Pt(10)

    doc.save(docx_path)
    print(f"Generated DOCX: {docx_path}")

if __name__ == '__main__':
    for role, name, npm, fname in STUDENTS:
        generate_md(role, name, npm, fname)
        generate_html(role, name, npm, fname)
        generate_docx(role, name, npm, fname)
