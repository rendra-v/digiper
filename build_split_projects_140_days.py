import datetime
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- 1. DATE CALCULATION FOR EXACTLY 140 DAYS (03/02/2026 s/d 03/08/2026) ---
start_date = datetime.date(2026, 2, 3)
end_date = datetime.date(2026, 8, 3)

holidays_user = {
    datetime.date(2026, 2, 16), # Imlek
    datetime.date(2026, 2, 17), # Cuti Bersama
    datetime.date(2026, 3, 18), # Cuti Bersama Nyepi
    datetime.date(2026, 3, 19), # Nyepi
    datetime.date(2026, 3, 20), # Cuti Bersama Idul Fitri
    datetime.date(2026, 3, 21), # Idul Fitri Hari 1
    datetime.date(2026, 3, 23), # Idul Fitri Hari 2
    datetime.date(2026, 3, 24), # Cuti Bersama Idul Fitri
    datetime.date(2026, 4, 3),  # Wafat Yesus Kristus
    datetime.date(2026, 5, 1),  # Hari Buruh
    datetime.date(2026, 5, 14), # Kenaikan Yesus Kristus
    datetime.date(2026, 5, 15), # Cuti Bersama Kenaikan
    datetime.date(2026, 5, 27), # Cuti Bersama Waisak
    datetime.date(2026, 5, 28), # Waisak
    datetime.date(2026, 6, 1),  # Hari Lahir Pancasila
    datetime.date(2026, 6, 16)  # Tahun Baru Islam / Idul Adha
}

day_names_id = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']

curr = start_date
working_days = [] # Exactly 140 days ending at 03/08/2026

while curr <= end_date:
    if curr.weekday() != 6 and curr not in holidays_user: # Exclude Sunday and 16 holidays
        day_name = day_names_id[curr.weekday()]
        date_str = curr.strftime('%d/%m/%Y')
        
        # WFH logic: All Saturdays WFH, alternating Fridays WFH
        if curr.weekday() == 5: # Sabtu
            work_mode = "WFH (Work From Home)"
        elif curr.weekday() == 4 and (len(working_days) % 2 == 1): # Alternate Fridays
            work_mode = "WFH (Work From Home)"
        else:
            work_mode = "WFO (Kantor MA)"
            
        working_days.append((day_name, date_str, work_mode, curr))
    curr += datetime.timedelta(days=1)

print(f"Generated {len(working_days)} effective working days (End: {working_days[-1][1]}).")

STUDENTS = [
    ("BE1", "Rifa Reza Fahlevi", "1402023060", "LOGBOOK_BE1_RIFA_REZA_FAHLEVI", "DAFTAR_HADIR_BE1_RIFA_REZA_FAHLEVI"),
    ("BE2", "Anggota Backend 2", "1402023061", "LOGBOOK_BE2", "DAFTAR_HADIR_BE2"),
    ("FE1", "Anggota Frontend 1", "1402023062", "LOGBOOK_FE1", "DAFTAR_HADIR_FE1"),
    ("FE2", "Anggota Frontend 2", "1402023063", "LOGBOOK_FE2", "DAFTAR_HADIR_FE2")
]

# --- REALISTIC TASKS SPLIT FOR 2 PROJECTS (70 DAYS ATKMA + 70 DAYS DIGIPER) ---
def get_task_for_day(role, day_idx):
    is_atk = day_idx <= 70 # Days 1-70 ATKMA, Days 71-140 DIGIPER
    
    if is_atk: # --- PROYEK 1: SISTEM INVENTARIS ATK (ATKMA) ---
        if role == "BE1":
            if day_idx == 1:
                return "[ATK] Analisis Masalah & Briefing Client: Identifikasi Kebutuhan ATK", "Wawancara dengan staf Mahkamah Agung & ketersediaan stok ATK", "Dokumentasi kebutuhan fitur terkumpul", "Istilah teknis barang instansi", "Menyesuaikan kamus istilah instansi", "Dokumentasi kebutuhan awal disepakati"
            elif day_idx <= 5:
                return f"[ATK] Analisis Masalah: Pembuatan FAI, Flowchart & ERD ATK Hari ke-{day_idx}", "Menyusun Formulir Analisis Informasi (FAI) & ERD database", "ERD master barang & transaksi disetujui", "Relasi kuota ruangan kompleks", "Kardinalitas One-to-Many Ruangan-Transaksi", "ERD Database ATK tervalidasi"
            elif day_idx <= 15:
                return f"[ATK] Backend: Setup Laravel 12, Migrasi & Model Master Barang Hari ke-{day_idx}", "Membuat migration & model Eloquent Barang, Ruangan & Category", "Schema DB Master Barang & Ruangan aktif", "Constraint min stock validation", "Constraint min:0 pada Eloquent", "Database schema bersih"
            elif day_idx <= 30:
                return f"[ATK] Backend: Stock Movement & Transaction Controller Hari ke-{day_idx}", "Membuat StockService & TransactionController (Permintaan Out/In)", "Fitur transaksi permintaan & auto reduce stok", "Race condition stok saat barang diminta bersamaan", "DB Transaction & lockForUpdate()", "Transaksi stok aman konsisten"
            elif day_idx <= 50:
                return f"[ATK] Backend: Stock Reconciliation & Audit Trail API Hari ke-{day_idx}", "Membuat Stock Reconciliation Controller & Audit Logs Logger", "Histori stok & rekonsiliasi barang aktif", "Stok fisik selisih dengan sistem", "Form Penyesuaian Stok + Catatan Alasan", "Audit log terverifikasi"
            elif day_idx <= 65:
                return f"[ATK] Backend: Laporan Kartu Stok / Buku Besar PDF (A4 Landscape) Hari ke-{day_idx}", "Membuat Export PDF Kartu Stok (Running Saldo, Saldo Awal, Tanda Tangan)", "PDF Laporan Kartu Stok A4 rapi presisi", "Page-break terpotong di tengah tabel", "Aturan CSS page-break-inside: avoid", "Laporan PDF Siap Cetak"
            else:
                return f"[ATK] Testing, Polish & Final Handover ATKMA Hari ke-{day_idx}", "Integration testing alur Permintaan -> Rekonsiliasi -> Laporan PDF", "Modul ATKMA 100% Pass Testing", "Deploy Railway env key mismatch", "Setting APP_KEY & DB SSL config", "Sistem ATKMA Live Railway Production"

        elif role == "BE2":
            if day_idx <= 10:
                return f"[ATK] Backend: Setup Authentication Fortify & Role Guard Hari ke-{day_idx}", "Membuat AuthController, Role Middleware (Admin/Pengawas)", "Sistem Auth & Permission Aktif", "Token Fortify CSRF mismatch", "Konfigurasi sanctum.stateful domains", "Auth Security Terjamin"
            elif day_idx <= 30:
                return f"[ATK] Backend: CRUD Master Ruangan & Kuota Permintaan Hari ke-{day_idx}", "Membuat RoomController, Validation Request & Room Policy", "CRUD Ruangan & Kuota aktif", "Kuota ruangan melebihi batas bulanan", "Custom Validation Rule RoomQuota", "Manajemen Ruangan Presisi"
            elif day_idx <= 50:
                return f"[ATK] Backend: User Management & API Response Standard Hari ke-{day_idx}", "Membuat UserController, Resource Formatter & Error Handler", "API User & Standardized JSON active", "Error 500 tidak informatif", "Global Exception Handler Handler.php", "API Error Response Clean"
            else:
                return f"[ATK] Backend: Redis Cache & Swagger API Docs ATKMA Hari ke-{day_idx}", "Memasang Redis caching master data & Swagger OpenAPI specs", "Dokumentasi API ATKMA lengkap", "Cache stale saat update barang", "Event Listener Cache Invalidation", "Modul BE2 ATKMA Complete"

        elif role == "FE1":
            if day_idx <= 15:
                return f"[ATK] Frontend: Perancangan UI/UX Figma & Component Library Hari ke-{day_idx}", "Membuat wireframe, shadcn/ui setup & Tailwind color MA", "Design System ATKMA siap pakai", "Komponen Modal belum accessible", "ARIA attributes & ESC handler", "Component library rapi"
            elif day_idx <= 35:
                return f"[ATK] Frontend: Halaman Dashboard & Data Table Stock Hari ke-{day_idx}", "Membuat Recharts Dashboard, Data Table & Filter Barang", "Dashboard statistik ATK interaktif", "Chart flickering saat filter", "React UseMemo & Skeleton State", "Dashboard UI Smooth"
            elif day_idx <= 55:
                return f"[ATK] Frontend: Form Permintaan Barang Dinamis Hari ke-{day_idx}", "Membuat Form Permintaan Multi-Item & Searchable Select", "Permintaan ATK multi-barang lancar", "Input kuota barang berulang", "Check duplicate item ID validator", "Form Transaksi Intuitif"
            else:
                return f"[ATK] Frontend: Responsive Polish & E2E Testing Playwright Hari ke-{day_idx}", "Polish tampilan mobile/tablet & E2E Playwright testing", "FE ATKMA 100% Pass E2E", "Tabel meluap di HP", "Wrapper overflow-x-auto", "Frontend ATKMA Complete"

        else: # FE2
            if day_idx <= 15:
                return f"[ATK] Frontend: Setup Inertia React, Types & Auth Page Hari ke-{day_idx}", "Setup Inertia.js, TypeScript interfaces & Login Form", "Halaman Login & Layout siap", "Input password tanpa toggle peek", "Eye Icon Toggle Password", "Auth UI Ramah User"
            elif day_idx <= 35:
                return f"[ATK] Frontend: CRUD Master Ruangan & User Admin UI Hari ke-{day_idx}", "Membuat UI Master Ruangan, Role Badge & Form Admin", "Management Ruangan UI aktif", "Pagination reset saat search", "Inertia preserveState & preserveScroll", "UI Navigation Lancar"
            elif day_idx <= 55:
                return f"[ATK] Frontend: Dark Mode & Export PDF/Excel Trigger UI Hari ke-{day_idx}", "Membuat Dark Mode toggle & Download Blob Progress Bar", "Export Laporan UI aktif", "Download Blob corrupt 0 byte", "Response type arraybuffer blob", "Download File Stabil"
            else:
                return f"[ATK] Frontend: UI Audit Trail & Finishing Touches Hari ke-{day_idx}", "Membuat UI Log Aktivitas, Toast Notifications & Polish", "Modul FE2 ATKMA Complete", "Toast muncul bertumpuk", "Single Toast Container Queue", "ATKMA App Production Ready"

    else: # --- PROYEK 2: SISTEM REKAPITULASI BIAYA PERKARA (DIGIPER) ---
        atk_idx = day_idx - 70 # 1 to 70 for DIGIPER
        if role == "BE1":
            if atk_idx <= 10:
                return f"[DIGIPER] Briefing Client & Engine Import Excel Rekap Perkara Hari ke-{atk_idx}", "Wawancara Tim Keuangan MA & parser Excel Cek/DPC/Tim Kep", "Engine import PhpSpreadsheet aktif", "Format tanggal Excel berupa angka serial", "Converter PhpSpreadsheet Date", "Excel Parser 100% Presisi"
            elif atk_idx <= 25:
                return f"[DIGIPER] Backend: Filter Otomatis Usia Perkara >= 90 Hari Hari ke-{atk_idx}", "Membuat query logic hitung tanggal putus - tanggal masuk >= 90", "Perkara >= 90 hari terfilter otomatis", "Perhitungan hari libur sabtu minggu", "Carbon diffInWeekdays()", "Filter 90 hari 100% akurat"
            elif atk_idx <= 45:
                return f"[DIGIPER] Backend: Business Logic Biaya Splitter & Honorarium Hari ke-{atk_idx}", "Membuat Kalkulator Biaya Perkara & Honor Majelis/PP/Ops", "Nominal Honorarium terhitung presisi", "Sisa pembagian desimal koma", "Pembulatan round() & alokasi sisa", "Total nominal 100% klop"
            elif atk_idx <= 60:
                return f"[DIGIPER] Backend: Multi-Worksheet Export Excel (Reconciliation Sheets) Hari ke-{atk_idx}", "Membuat Export Multi-Sheet (Cek, DPC, Tim Kep, Pemilah, Op Staf)", "Excel Rekapitulasi Multi-Sheet siap", "Formula Excel SUM corrupt", "Writing raw Excel formula =SUM()", "Multi-Sheet Excel Perfect"
            else:
                return f"[DIGIPER] Final Testing, Security Audit & Production Deployment Hari ke-{atk_idx}", "Full SIT & UAT bersama Staf MA & Deployment Server MA", "DIGIPER Live Production MA RI", "Memory limit saat export 10.000 row", "Eloquent Chunk & Stream Download", "DIGIPER App Complete Live"

        elif role == "BE2":
            if atk_idx <= 15:
                return f"[DIGIPER] Backend: Master 8 Jenis Perkara & Senioritas Hakim Hari ke-{atk_idx}", "Membuat Master Perkara (Perdata, Pidana, dll) & Senioritas", "Master Data Perkara & Hakim aktif", "Gelar Hakim Dr. S.H. M.H. mengganggu matching", "Regex Stripping Gelar Akademik", "Fuzzy Matching 98% Presisi"
            elif atk_idx <= 40:
                return f"[DIGIPER] Backend: API Audit Trail & Rollback Session Import Hari ke-{atk_idx}", "Membuat Session Import Handler & Rollback Transaction", "Fitur Batal/Rollback Import aktif", "Rollback meninggalkan data gantung", "DB Transaction Rollback All", "Rollback Session Aman"
            else:
                return f"[DIGIPER] Backend: Testing SIT, Performance & Load Test Hari ke-{atk_idx}", "Testing beban data 50.000 perkara & API Optimization", "DIGIPER Backend Pass SIT 100%", "Response time import > 10 detik", "Queued Jobs Horizon Background Processing", "DIGIPER Backend Super Fast"

        elif role == "FE1":
            if atk_idx <= 20:
                return f"[DIGIPER] Frontend: Drag & Drop Excel Uploader & Progress Bar Hari ke-{atk_idx}", "Membuat Drag & Drop Zone, File Validation & Upload Bar", "Uploader Excel interaktif", "Browser freeze saat upload file besar", "Web Worker File Reader", "Upload UX Smooth"
            elif atk_idx <= 45:
                return f"[DIGIPER] Frontend: Interactive Preview Table 90 Hari & Filters Hari ke-{atk_idx}", "Membuat Table Preview Perkara Putus & Advance Filter", "Preview Data Perkara ciamik", "Tabel 5.000 row lambat scroll", "Tanstack Virtual Table Windowing", "Table Render Instan"
            else:
                return f"[DIGIPER] Frontend: Live Preview PDF Rekapitulasi & Final UAT Hari ke-{atk_idx}", "Membuat PDF Live Viewer (pdfjs) & UAT bersama Pembimbing", "DIGIPER FE Pass UAT 100%", "PDF Viewer terblokir popup blocker", "Render Canvas PDF Inline", "DIGIPER Frontend Ready"

        else: # FE2
            if atk_idx <= 20:
                return f"[DIGIPER] Frontend: Reorder Drag & Drop Senioritas Hakim Agung Hari ke-{atk_idx}", "Membuat Drag-and-Drop Reorder list Senioritas Hakim", "Reorder Senioritas Hakim instan", "Urutan kembali ke awal saat refresh", "Sync order state to Backend API", "Reorder State Persistent"
            elif atk_idx <= 45:
                return f"[DIGIPER] Frontend: Dashboard Honorarium & Breakdown Modal Hari ke-{atk_idx}", "Membuat Cards Summary Honor & Modal Detail Rincian Biaya", "Dashboard Honorarium terwujud", "Nominal angka tanpa format Rp", "Intl.NumberFormat Rupiah Helper", "Rincian Keuangan Jelas"
            else:
                return f"[DIGIPER] Frontend: Modal Rollback Safety Confirmation & UAT Polish Hari ke-{atk_idx}", "Membuat Modal Konfirmasi Rollback (Type 'HAPUS') & Polish", "DIGIPER FE Complete Pass UAT", "Klik Hapus tidak sengaja", "Type Confirmation Keyword", "DIGIPER App 100% Ready"

# --- DOCX HELPERS ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="1F497D", sz="4"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            borders.append(border)
        tblPr[0].append(borders)

# --- 2. GENERATE LOGBOOK FILES (2 PROJECTS) ---
def generate_logbook_files(role, name, npm, logbook_fname):
    md_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.md"
    html_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.html"
    docx_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.docx"

    # Markdown
    md_content = f"# LOGBOOK LAPORAN KEGIATAN MAGANG ({role})\n"
    md_content += f"**Nama Mahasiswa:** {name}  \n"
    md_content += f"**NPM:** {npm}  \n"
    md_content += f"**Instansi:** Mahkamah Agung Republik Indonesia  \n"
    md_content += f"**Periode Magang:** 03 Februari 2026 – {working_days[-1][1]} (140 Hari Kerja / 1.120 Jam)  \n"
    md_content += f"**Pembagian Proyek:**  \n"
    md_content += f"- 📦 **PROYEK 1 (ATKMA):** Hari ke-1 s.d. 70 (03/02/2026 – 13/05/2026 | 560 Jam) — Sistem Inventaris ATK  \n"
    md_content += f"- ⚖️ **PROYEK 2 (DIGIPER):** Hari ke-71 s.d. 140 (16/05/2026 – 03/08/2026 | 560 Jam) — Sistem Rekapitulasi Biaya Perkara  \n\n"
    md_content += "---\n\n"
    md_content += "| NO | HARI/TANGGAL | RENCANA | REALISASI | HASIL | PROBLEM | SOLUSI | HASIL AKHIR |\n"
    md_content += "|:--:|:---:|:---|:---|:---|:---|:---|:---|\n"

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        md_content += f"| {idx} | {day_name}, {date_str} | {rencana} | {realisasi} | {hasil} | {problem} | {solusi} | {hasil_akhir} |\n"

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    # HTML
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>LOGBOOK MAGANG - {name} ({role})</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 30px; background-color: #f8f9fa; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1 {{ color: #1f497d; margin-bottom: 5px; font-size: 22px; }}
        .info {{ margin-bottom: 20px; font-size: 14px; line-height: 1.6; background: #eef2f7; padding: 15px; border-left: 4px solid #1f497d; border-radius: 4px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 12px; }}
        th {{ background-color: #1f497d; color: white; padding: 8px; text-align: center; border: 1px solid #1f497d; font-weight: bold; }}
        td {{ padding: 7px 8px; border: 1px solid #dce4ec; vertical-alignment: top; }}
        tr:nth-child(even) {{ background-color: #f2f5f9; }}
        .center {{ text-align: center; }}
        .badge-atk {{ background-color: #0284c7; color: white; padding: 2px 6px; border-radius: 4px; font-size: 11px; font-weight: bold; }}
        .badge-digiper {{ background-color: #15803d; color: white; padding: 2px 6px; border-radius: 4px; font-size: 11px; font-weight: bold; }}
    </style>
</head>
<body>
<div class="container">
    <h1>LOGBOOK LAPORAN KEGIATAN MAGANG — 2 PROYEK (ATKMA &amp; DIGIPER)</h1>
    <div class="info">
        <b>Nama Mahasiswa:</b> {name}<br>
        <b>NPM:</b> {npm}<br>
        <b>Peran / Posisi:</b> {role}<br>
        <b>Instansi:</b> Mahkamah Agung Republik Indonesia<br>
        <b>Periode Magang:</b> 03 Februari 2026 – {working_days[-1][1]} (140 Hari Kerja / 1.120 Jam Total)<br>
        <b>Pembagian Proyek:</b><br>
        &bull; <span class="badge-atk">PROYEK 1: ATKMA</span> Hari ke-1 s.d. 70 (03/02/2026 – 13/05/2026 | 560 Jam) — Sistem Inventaris ATK Mahkamah Agung<br>
        &bull; <span class="badge-digiper">PROYEK 2: DIGIPER</span> Hari ke-71 s.d. 140 (16/05/2026 – 03/08/2026 | 560 Jam) — Sistem Rekapitulasi Biaya Perkara Mahkamah Agung
    </div>
    <table>
        <thead>
            <tr>
                <th width="4%">NO</th>
                <th width="14%">HARI/TANGGAL</th>
                <th width="16%">RENCANA</th>
                <th width="18%">REALISASI</th>
                <th width="16%">HASIL</th>
                <th width="11%">PROBLEM</th>
                <th width="11%">SOLUSI</th>
                <th width="10%">HASIL AKHIR</th>
            </tr>
        </thead>
        <tbody>
"""
    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        html_content += f"""            <tr>
                <td class="center">{idx}</td>
                <td class="center">{day_name}, {date_str}</td>
                <td>{rencana}</td>
                <td>{realisasi}</td>
                <td>{hasil}</td>
                <td>{problem}</td>
                <td>{solusi}</td>
                <td>{hasil_akhir}</td>
            </tr>\n"""

    html_content += "        </tbody>\n    </table>\n</div>\n</body>\n</html>"
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    # DOCX
    doc = docx.Document()
    sec = doc.sections[0]
    sec.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11.69)
    sec.page_height = Inches(8.27)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    p = doc.add_paragraph()
    r = p.add_run(f"LOGBOOK LAPORAN KEGIATAN MAGANG ({role})\n")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(8)
    p_info.add_run(f"Nama: {name} | NPM: {npm} | Instansi: Mahkamah Agung RI\nPeriode: 03 Februari 2026 – {working_days[-1][1]} | Total Hari Kerja: 140 Hari (1.120 Jam Total)\nProyek 1 (ATKMA): Hari ke-1 s.d 70 (560 Jam) | Proyek 2 (DIGIPER): Hari ke-71 s.d 140 (560 Jam)")

    table = doc.add_table(rows=len(working_days) + 1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="1F497D")

    headers = ["NO", "HARI/TANGGAL", "RENCANA", "REALISASI", "HASIL", "PROBLEM", "SOLUSI", "HASIL AKHIR"]
    col_widths = [Inches(0.4), Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.7), Inches(1.3), Inches(1.3), Inches(1.2)]

    for c_idx, h_text in enumerate(headers):
        cell = table.cell(0, c_idx)
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.width = col_widths[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        vals = [str(idx), f"{day_name}, {date_str}", rencana, realisasi, hasil, problem, solusi, hasil_akhir]
        bg_color = "F2F5F9" if idx % 2 == 1 else "FFFFFF"

        for c_idx, val in enumerate(vals):
            cell = table.cell(idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            if c_idx in [0, 1]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(val)
            run.font.size = Pt(8.5)
            
            if "[ATK]" in val:
                run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7) # Blue ATK
            elif "[DIGIPER]" in val:
                run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D) # Green DIGIPER

    doc.save(docx_path)
    print(f"Generated Logbook Split (140 Days): {docx_path}")

# --- 3. GENERATE DAFTAR HADIR FILES FOR 140 DAYS ---
def generate_hadir_files(role, name, npm, hadir_fname):
    md_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.md"
    html_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.html"
    docx_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.docx"

    # Markdown
    md_content = f"# DAFTAR HADIR MAGANG MAHASISWA\n"
    md_content += f"**Tahun Ajaran:** 2025/2026  \n"
    md_content += f"**Instansi:** Mahkamah Agung Republik Indonesia  \n"
    md_content += f"**Nama Mahasiswa:** {name}  \n"
    md_content += f"**NPM:** {npm}  \n"
    md_content += f"**Peran:** {role}  \n"
    md_content += f"**Total Kehadiran:** 140 Hari Hadir (100% Hadir / 1.120 Jam)  \n"
    md_content += f"**Pelaksanaan Proyek:** ATKMA (Hari 1–70) & DIGIPER (Hari 71–140)  \n\n"
    md_content += "---\n\n"
    md_content += "| NO | NAMA MAHASISWA | NPM | TANGGAL HADIR | JUMLAH: H | JUMLAH: I | JUMLAH: S | JUMLAH: A | KETERANGAN |\n"
    md_content += "|:--:|:---|:---:|:---:|:--:|:--:|:--:|:--:|:---|\n"

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        md_content += f"| {idx} | {name} | {npm} | {day_name}, {date_str} | ✓ | | | | {work_mode} |\n"

    md_content += "\n---\n"
    md_content += "**Keterangan:**\n"
    md_content += "- **H**: Hadir (✓ Checkmark - 140 Hari)\n"
    md_content += "- **I**: Izin (0)\n"
    md_content += "- **S**: Sakit (0)\n"
    md_content += "- **A**: Alpa (0)\n\n"
    md_content += "**Catatan Kehadiran:**\n"
    md_content += "- **WFO (Work From Office)**: Hari Senin s.d. Kamis & Jumat Terjadwal (Kantor MA)\n"
    md_content += "- **WFH (Work From Home)**: Hari Sabtu & Jumat Terjadwal\n\n"
    md_content += f"Jakarta, {working_days[-1][1]}  \n"
    md_content += "**Instruktur / Pembimbing Magang Instansi:**  \n\n\n"
    md_content += "__________________________________________  \n"
    md_content += "NIP. …………………………………………  \n"

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    # HTML
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>DAFTAR HADIR MAGANG - {name}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 30px; background-color: #f8f9fa; }}
        .container {{ max-width: 1100px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #1f497d; margin-bottom: 5px; font-size: 22px; text-decoration: underline; }}
        h2 {{ text-align: center; color: #555; margin-top: 0; font-size: 15px; font-weight: normal; margin-bottom: 25px; }}
        .info {{ margin-bottom: 20px; font-size: 14px; line-height: 1.6; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 13px; }}
        th {{ background-color: #1f497d; color: white; padding: 8px; text-align: center; border: 1px solid #1f497d; font-weight: bold; }}
        td {{ padding: 7px 8px; border: 1px solid #dce4ec; text-align: center; }}
        td.left {{ text-align: left; }}
        tr:nth-child(even) {{ background-color: #f2f5f9; }}
        .wfh {{ color: #d97706; font-weight: bold; }}
        .wfo {{ color: #15803d; font-weight: bold; }}
        .check {{ color: #15803d; font-weight: bold; font-size: 16px; }}
        .footer {{ margin-top: 30px; float: right; text-align: center; width: 280px; font-size: 14px; }}
        .clear {{ clear: both; }}
    </style>
</head>
<body>
<div class="container">
    <h1>DAFTAR HADIR MAGANG</h1>
    <h2>Tahun Ajaran 2025/2026 — Mahkamah Agung Republik Indonesia</h2>
    <div class="info">
        <b>Nama Mahasiswa:</b> {name}<br>
        <b>NPM:</b> {npm}<br>
        <b>Peran / Posisi:</b> {role}<br>
        <b>Total Kehadiran:</b> 140 Hari Hadir (100% Hadir / 1.120 Jam)<br>
        <b>Pelaksanaan Proyek:</b> Proyek 1 ATKMA (Hari 1–70) &amp; Proyek 2 DIGIPER (Hari 71–140)
    </div>
    <table>
        <thead>
            <tr>
                <th rowspan="2" width="5%">NO</th>
                <th rowspan="2" width="25%">NAMA MAHASISWA</th>
                <th rowspan="2" width="13%">NPM</th>
                <th rowspan="2" width="22%">TANGGAL HADIR</th>
                <th colspan="4" width="16%">JUMLAH</th>
                <th rowspan="2" width="19%">KETERANGAN</th>
            </tr>
            <tr>
                <th width="4%">H</th>
                <th width="4%">I</th>
                <th width="4%">S</th>
                <th width="4%">A</th>
            </tr>
        </thead>
        <tbody>
"""
    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        cls = "wfh" if "WFH" in work_mode else "wfo"
        html_content += f"""            <tr>
                <td>{idx}</td>
                <td class="left">{name}</td>
                <td>{npm}</td>
                <td>{day_name}, {date_str}</td>
                <td class="check">✓</td>
                <td></td>
                <td></td>
                <td></td>
                <td class="{cls}">{work_mode}</td>
            </tr>\n"""

    html_content += f"""        </tbody>
    </table>
    <div style="margin-top: 20px; font-size: 13px; line-height: 1.6;">
        H : Hadir (✓ Checkmark - 140 Hari)<br>
        I : Izin (0)<br>
        S : Sakit (0)<br>
        A : Alpa (0)
    </div>
    <div class="footer">
        Jakarta, {working_days[-1][1]}<br>
        <b>Nama Instansi Magang: Mahkamah Agung RI</b><br><br><br><br>
        __________________________________________<br>
        <b>Instruktur Magang</b>
    </div>
    <div class="clear"></div>
</div>
</body>
</html>"""

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    # DOCX
    doc = docx.Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR HADIR MAGANG\n")
    r.bold = True
    r.underline = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    r2 = p.add_run("Tahun Ajaran 2025/2026 — Mahkamah Agung Republik Indonesia\n")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(8)
    p_meta.add_run(f"Nama Mahasiswa : {name}\nNPM             : {npm}\nPeran / Posisi  : {role}\nTotal Kehadiran : 140 Hari Hadir (100% Hadir / 1.120 Jam)")

    table = doc.add_table(rows=len(working_days) + 2, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="1F497D")

    col_widths = [Inches(0.4), Inches(1.8), Inches(1.0), Inches(1.4), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(1.6)]

    hdr_cells_r0 = table.rows[0].cells
    hdr_cells_r1 = table.rows[1].cells

    hdr_cells_r0[0].text = "NO"
    hdr_cells_r0[0].merge(hdr_cells_r1[0])
    
    hdr_cells_r0[1].text = "NAMA MAHASISWA"
    hdr_cells_r0[1].merge(hdr_cells_r1[1])

    hdr_cells_r0[2].text = "NPM"
    hdr_cells_r0[2].merge(hdr_cells_r1[2])

    hdr_cells_r0[3].text = "TANGGAL HADIR"
    hdr_cells_r0[3].merge(hdr_cells_r1[3])

    hdr_cells_r0[4].text = "JUMLAH"
    hdr_cells_r0[4].merge(hdr_cells_r0[5]).merge(hdr_cells_r0[6]).merge(hdr_cells_r0[7])

    hdr_cells_r1[4].text = "H"
    hdr_cells_r1[5].text = "I"
    hdr_cells_r1[6].text = "S"
    hdr_cells_r1[7].text = "A"

    hdr_cells_r0[8].text = "KETERANGAN"
    hdr_cells_r0[8].merge(hdr_cells_r1[8])

    for row in [table.rows[0], table.rows[1]]:
        for cell in row.cells:
            set_cell_background(cell, "1F497D")
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9)

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        r_idx = idx + 1
        row_cells = table.rows[r_idx].cells
        bg_color = "F2F5F9" if idx % 2 == 1 else "FFFFFF"

        vals = [str(idx), name, npm, f"{day_name}, {date_str}", "✓", "", "", "", work_mode]

        for c_idx, val in enumerate(vals):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = col_widths[c_idx]

            p = cell.paragraphs[0]
            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(val)
            run.font.size = Pt(8.5)

            if c_idx == 4:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
            elif "WFH" in val:
                run.bold = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    doc.add_paragraph()
    p_leg = doc.add_paragraph()
    p_leg.paragraph_format.space_after = Pt(15)
    r_leg = p_leg.add_run("H  : Hadir (✓ Checkmark - 140 Hari)\nI   : Izin\nS   : Sakit\nA  : Alpa")
    r_leg.font.size = Pt(9.5)

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.paragraph_format.space_before = Pt(15)
    r_sig = p_sig.add_run(f"Jakarta, {working_days[-1][1]}\nNama instansi magang: Mahkamah Agung RI\n\n\n\n__________________________________________\nInstruktur Magang")
    r_sig.font.size = Pt(10)

    doc.save(docx_path)
    print(f"Generated Hadir DOCX (140 Days): {docx_path}")

if __name__ == '__main__':
    for role, name, npm, logbook_fname, hadir_fname in STUDENTS:
        generate_logbook_files(role, name, npm, logbook_fname)
        generate_hadir_files(role, name, npm, hadir_fname)
