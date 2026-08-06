import datetime
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- 1. DATE CALCULATION & HOLIDAY FILTERING ---
start_date = datetime.date(2026, 2, 3)
holidays_user = {
    datetime.date(2026, 2, 16), # Imlek
    datetime.date(2026, 2, 17), # Cuti Bersama
    datetime.date(2026, 3, 18), # Cuti Bersama Nyepi
    datetime.date(2026, 3, 19), # Nyepi
    datetime.date(2026, 3, 20), # Cuti Bersama Idul Fitri
    datetime.date(2026, 3, 21), # Idul Fitri Hari 1
    datetime.date(2026, 3, 23), # Idul Fitri Hari 2
    datetime.date(2026, 3, 24), # Cuti Bersama Idul Fitri
    datetime.date(2026, 4, 3),  # Wafat Yesus Kristus
    datetime.date(2026, 5, 1),  # Hari Buruh
    datetime.date(2026, 5, 14), # Kenaikan Yesus Kristus
    datetime.date(2026, 5, 15), # Cuti Bersama Kenaikan
    datetime.date(2026, 5, 27), # Cuti Bersama Waisak
    datetime.date(2026, 5, 28), # Waisak
    datetime.date(2026, 6, 1),  # Hari Lahir Pancasila
    datetime.date(2026, 6, 16)  # Tahun Baru Islam / Idul Adha
}

day_names_id = ['Senin', 'Selasa', 'Rabu', 'Kamis', 'Jumat', 'Sabtu', 'Minggu']

curr = start_date
working_days = [] # Exactly 144 days

while len(working_days) < 144:
    if curr.weekday() != 6 and curr not in holidays_user: # Exclude Sunday and 16 holidays
        day_name = day_names_id[curr.weekday()]
        date_str = curr.strftime('%d/%m/%Y')
        
        # WFH logic: All Saturdays WFH, alternating Fridays WFH
        if curr.weekday() == 5: # Sabtu
            work_mode = "WFH (Work From Home)"
        elif curr.weekday() == 4 and (len(working_days) % 2 == 1): # Alternate Fridays
            work_mode = "WFH (Work From Home)"
        else:
            work_mode = "WFO (Kantor MA)"
            
        working_days.append((day_name, date_str, work_mode, curr))
    curr += datetime.timedelta(days=1)

STUDENTS = [
    ("BE1", "Rifa Reza Fahlevi", "1402023060", "LOGBOOK_BE1_RIFA_REZA_FAHLEVI", "DAFTAR_HADIR_BE1_RIFA_REZA_FAHLEVI"),
    ("BE2", "Anggota Backend 2", "1402023061", "LOGBOOK_BE2", "DAFTAR_HADIR_BE2"),
    ("FE1", "Anggota Frontend 1", "1402023062", "LOGBOOK_FE1", "DAFTAR_HADIR_FE1"),
    ("FE2", "Anggota Frontend 2", "1402023063", "LOGBOOK_FE2", "DAFTAR_HADIR_FE2")
]

def get_task_for_day(role, day_idx):
    is_atk = day_idx <= 72
    proj_name = "ATK" if is_atk else "DIGIPER"
    
    if role == "BE1":
        if day_idx == 1:
            return "Analisis Masalah + Briefing Rapat: Wawancara Client & Kebutuhan ATK", "Wawancara dengan staf Mahkamah Agung & identifikasi kebutuhan", "Catatan kebutuhan fitur terkumpul", "Istilah teknis barang instansi", "Menyesuaikan kamus istilah instansi", "Dokumentasi kebutuhan awal disepakati"
        elif day_idx <= 10:
            return f"Pengembangan Schema Database & Core API {proj_name} Hari ke-{day_idx}", f"Membuat migration & controller master {proj_name}", f"Endpoint CRUD {proj_name} aktif", "Validasi stok minimum 0", "Menambahkan constraint validation min:0", "API Master tervalidasi"
        elif day_idx <= 25:
            return f"Pengembangan Service Layer & Business Logic {proj_name} Hari ke-{day_idx}", f"Membuat logic pengajuan & approval {proj_name}", f"Workflow approval {proj_name} selesai", "Konflik status saat diajukan bersamaan", "Database transaction & lockForUpdate", "Keamanan transaksi terjamin"
        elif day_idx <= 45:
            return f"Pengembangan Modul Laporan & Export PDF/Excel {proj_name} Hari ke-{day_idx}", f"Membuat controller rekapitulasi data & export", f"Laporan {proj_name} dapat dicetak", "Layout PDF berantakan", "Mengatur page-break CSS landscape", "PDF Laporan rapi presisi"
        elif day_idx <= 72:
            return f"Security Hardening, Optimization & Testing {proj_name} Hari ke-{day_idx}", f"Menerapkan rate limiting & audit log {proj_name}", f"Sistem {proj_name} kebal serangan", "Memory limit PHP saat load besar", "Menggunakan chunk processing Eloquent", "Performa backend optimal"
        elif day_idx <= 100:
            return f"Pengembangan Backend DIGIPER: Excel Import & Auto Filter 90 Hari Hari ke-{day_idx}", f"Membuat parser Excel & logic usia perkara >= 90 hari", f"Perkara >= 90 hari terfilter otomatis", "Format tanggal Excel berupa angka serial", "Converter PhpSpreadsheet Date", "Filter 90 hari 100% akurat"
        elif day_idx <= 130:
            return f"Pengembangan Backend DIGIPER: Biaya Splitter & Honorarium Hakim Hari ke-{day_idx}", f"Membuat logic pemecahan biaya & honor Majelis/PP/Ops", f"Nominal honorarium terhitung presisi", "Sisa pembagian desimal koma", "Pembulatan round() & alokasi sisa", "Total pecahan 100% klop"
        else:
            return f"Finalization, Split Export & Deployment DIGIPER Hari ke-{day_idx}", f"Membuat class Multi-Worksheet Export & deployment prod", f"DIGIPER live di server Mahkamah Agung", "SSL Config Nginx web server", "Konfigurasi SSL HTTPS secure", "DIGIPER Live Production"

    elif role == "BE2":
        if day_idx <= 15:
            return f"Pengembangan Auth Sanctum & Master Data {proj_name} Hari ke-{day_idx}", f"Membuat AuthController, JWT/Sanctum & User Management", f"Auth & Master Data API aktif", "Token expired mendadak", "Menyesuaikan TTL token di config", "Auth API stabil"
        elif day_idx <= 35:
            return f"Pengembangan API Stock Movement & Audit Logs {proj_name} Hari ke-{day_idx}", f"Membuat query histori pergerakan & activity tracker", f"Histori pergerakan barang terekam", "Query histori lambat data besar", "Indexing kolom item_id & created_at", "Query histori instan"
        elif day_idx <= 72:
            return f"API Caching, Rate Limiting & Swagger Specs {proj_name} Hari ke-{day_idx}", f"Memasang Redis cache & Swagger OpenAPI documentation", f"Dokumentasi API {proj_name} terpublikasi", "Cache stale saat edit data", "Cache Invalidation event listener", "Cache selalu akurat"
        elif day_idx <= 110:
            return f"Pengembangan Backend DIGIPER: Master Jenis & Fuzzy Matching Hari ke-{day_idx}", f"Membuat master 8 jenis perkara & fuzzy matching Hakim", f"Fuzzy matching nama Hakim 98% akurat", "Gelar Dr. S.H. M.H. mengganggu match", "Regex stripping gelar Hakim", "Matching rate 98% presisi"
        else:
            return f"Backend DIGIPER: Master Pejabat, Audit Trail & SIT Hari ke-{day_idx}", f"Membuat master Pejabat, audit trail & testing SIT", f"Module BE2 DIGIPER Pass SIT", "Toleransi selisih koma pembulatan", "Toleransi selisih nominal < Rp 10", "Keuangan terverifikasi"

    elif role == "FE1":
        if day_idx <= 20:
            return f"Perancangan UI/UX & Component Library {proj_name} Hari ke-{day_idx}", f"Membuat wireframe Figma, Design System & Base Components", f"Component library {proj_name} terbentuk", "Modal belum menangani key ESC", "Event listener keydown Escape", "Modal UI aksesibel"
        elif day_idx <= 45:
            return f"Pengembangan Dashboard Charts & Data Table {proj_name} Hari ke-{day_idx}", f"Membuat komponen Recharts & DataTable responsive", f"Visualisasi Dashboard ciamik", "Chart flickering saat re-fetch", "Loading skeleton condition", "Transisi data chart mulus"
        elif day_idx <= 72:
            return f"Responsiveness, A11y & E2E Testing Playwright {proj_name} Hari ke-{day_idx}", f"Polish mobile layout, ARIA labels & E2E Playwright test", f"Aplikasi responsive & ramah difabel", "Tabel meluap di HP kecil", "Wrapper overflow-x-auto", "Tampilan mobile rapi"
        elif day_idx <= 110:
            return f"Pengembangan Frontend DIGIPER: Upload Drag & Drop & Preview Hari ke-{day_idx}", f"Membuat ExcelUploader zone, progress bar & preview table", f"Drag & Drop Upload Excel interaktif", "File besar browser hang", "Memindahkan read file ke Web Worker", "Upload Excel lancar"
        else:
            return f"Frontend DIGIPER: Modal Detail, Live Preview PDF & SIT Hari ke-{day_idx}", f"Membuat Modal Detail Rekapan, pdfjs Live Preview & SIT", f"Frontend DIGIPER Complete SIT Pass", "Viewer PDF terblokir popup", "pdfjs-dist render canvas", "Live Preview PDF Terjamin"

    else:
        if day_idx <= 20:
            return f"Pengembangan Form UI & Auth System {proj_name} Hari ke-{day_idx}", f"Membuat Form Login, Form Permintaan Dinamis & History UI", f"Form UI & Auth interaktif", "Input password tanpa caps warning", "Warning indicator Caps Lock", "UI Auth ramah"
        elif day_idx <= 45:
            return f"Pengembangan Master Ruangan & Download Blob Handler {proj_name} Hari ke-{day_idx}", f"Membuat UI Master Ruangan, Dark Mode & Blob Handler", f"Download Laporan & Dark Mode aktif", "File download corrupt error 500", "Check response type JSON", "Download file stabil"
        elif day_idx <= 72:
            return f"State Management Zustand & Performance Tuning {proj_name} Hari ke-{day_idx}", f"Implementasi Zustand store, Lazy loading & Image compression", f"Bundle size < 300KB fast load", "Re-render komponen berlebihan", "Selector spesifik Zustand", "Render UI efisien"
        elif day_idx <= 110:
            return f"Pengembangan Frontend DIGIPER: Master Hakim Senioritas Reorder Hari ke-{day_idx}", f"Membuat Drag-and-Drop Reorder Senioritas Hakim Agung", f"Senioritas Hakim terurut instan", "Input NIP tanpa regex format", "Regex validation NIP 18 digit", "Validation NIP presisi"
        else:
            return f"Frontend DIGIPER: Dashboard Honorarium, Rollback UI & UAT Hari ke-{day_idx}", f"Membuat UI Dashboard Honor, Modal Rollback Import & UAT", f"100% UAT Pass bersama Staf MA", "Rollback terklik tidak sengaja", "Konfirmasi ketik kata HAPUS", "Rollback UI Safe"

# --- DOCX FORMATTING HELPERS ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="1F497D", sz="4"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            borders.append(border)
        tblPr[0].append(borders)

# --- 2. GENERATE LOGBOOK FILES ---
def generate_logbook_files(role, name, npm, logbook_fname):
    md_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.md"
    html_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.html"
    docx_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{logbook_fname}.docx"

    # Markdown
    md_content = f"# LOGBOOK LAPORAN KEGIATAN MAGANG ({role})\n"
    md_content += f"**Nama Mahasiswa:** {name}  \n"
    md_content += f"**NPM:** {npm}  \n"
    md_content += f"**Instansi:** Mahkamah Agung Republik Indonesia  \n"
    md_content += f"**Periode Magang:** 03 Februari 2026 – {working_days[-1][1]}  \n"
    md_content += f"**Total Hari Kerja Efektif:** 144 Hari (Jam kerja: 8 jam/hari)  \n"
    md_content += f"**Total Jam Magang:** 1.152 Jam (Lulus syarat minimal 1.035 jam, kelebihan 117 jam)  \n"
    md_content += f"**Catatan Libur:** 16 Hari Libur Nasional / Cuti Bersama pada hari kerja tidak dihitung.  \n\n"
    md_content += "---\n\n"
    md_content += "| NO | HARI/TANGGAL | RENCANA | REALISASI | HASIL | PROBLEM | SOLUSI | HASIL AKHIR |\n"
    md_content += "|:--:|:---:|:---|:---|:---|:---|:---|:---|\n"

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        md_content += f"| {idx} | {day_name}, {date_str} | {rencana} | {realisasi} | {hasil} | {problem} | {solusi} | {hasil_akhir} |\n"

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    # HTML
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>LOGBOOK MAGANG - {name} ({role})</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 30px; background-color: #f8f9fa; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1 {{ color: #1f497d; margin-bottom: 5px; font-size: 22px; }}
        .info {{ margin-bottom: 20px; font-size: 14px; line-height: 1.6; background: #eef2f7; padding: 15px; border-left: 4px solid #1f497d; border-radius: 4px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 12px; }}
        th {{ background-color: #1f497d; color: white; padding: 8px; text-align: center; border: 1px solid #1f497d; font-weight: bold; }}
        td {{ padding: 7px 8px; border: 1px solid #dce4ec; vertical-alignment: top; }}
        tr:nth-child(even) {{ background-color: #f2f5f9; }}
        .center {{ text-align: center; }}
    </style>
</head>
<body>
<div class="container">
    <h1>LOGBOOK LAPORAN KEGIATAN MAGANG — 144 HARI KERJA EFEKTIF (1.152 JAM)</h1>
    <div class="info">
        <b>Nama Mahasiswa:</b> {name}<br>
        <b>NPM:</b> {npm}<br>
        <b>Peran / Posisi:</b> {role}<br>
        <b>Instansi:</b> Mahkamah Agung Republik Indonesia<br>
        <b>Periode Magang:</b> 03 Februari 2026 – {working_days[-1][1]}<br>
        <b>Total Hari Kerja:</b> 144 Hari Kerja Efektif &times; 8 Jam = <b>1.152 Jam</b> (Syarat Minimal: 1.035 Jam | Kelebihan: 117 Jam)<br>
        <b>Keterangan Libur:</b> 16 Hari Libur Nasional &amp; Cuti Bersama pada hari kerja dikeluarkan dari perhitungan.
    </div>
    <table>
        <thead>
            <tr>
                <th width="4%">NO</th>
                <th width="14%">HARI/TANGGAL</th>
                <th width="16%">RENCANA</th>
                <th width="18%">REALISASI</th>
                <th width="16%">HASIL</th>
                <th width="11%">PROBLEM</th>
                <th width="11%">SOLUSI</th>
                <th width="10%">HASIL AKHIR</th>
            </tr>
        </thead>
        <tbody>
"""
    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        html_content += f"""            <tr>
                <td class="center">{idx}</td>
                <td class="center">{day_name}, {date_str}</td>
                <td>{rencana}</td>
                <td>{realisasi}</td>
                <td>{hasil}</td>
                <td>{problem}</td>
                <td>{solusi}</td>
                <td>{hasil_akhir}</td>
            </tr>\n"""

    html_content += "        </tbody>\n    </table>\n</div>\n</body>\n</html>"
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    # DOCX
    doc = docx.Document()
    sec = doc.sections[0]
    sec.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11.69)
    sec.page_height = Inches(8.27)
    sec.left_margin = Inches(0.5)
    sec.right_margin = Inches(0.5)
    sec.top_margin = Inches(0.5)
    sec.bottom_margin = Inches(0.5)

    p = doc.add_paragraph()
    r = p.add_run(f"LOGBOOK LAPORAN KEGIATAN MAGANG ({role})\n")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p_info = doc.add_paragraph()
    p_info.paragraph_format.space_after = Pt(8)
    p_info.add_run(f"Nama: {name} | NPM: {npm} | Instansi: Mahkamah Agung RI\nPeriode: 03 Februari 2026 – {working_days[-1][1]} | Total Hari Kerja: 144 Hari (1.152 Jam Total)")

    table = doc.add_table(rows=len(working_days) + 1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="1F497D")

    headers = ["NO", "HARI/TANGGAL", "RENCANA", "REALISASI", "HASIL", "PROBLEM", "SOLUSI", "HASIL AKHIR"]
    col_widths = [Inches(0.4), Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.7), Inches(1.3), Inches(1.3), Inches(1.2)]

    for c_idx, h_text in enumerate(headers):
        cell = table.cell(0, c_idx)
        set_cell_background(cell, "1F497D")
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        cell.width = col_widths[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        rencana, realisasi, hasil, problem, solusi, hasil_akhir = get_task_for_day(role, idx)
        vals = [str(idx), f"{day_name}, {date_str}", rencana, realisasi, hasil, problem, solusi, hasil_akhir]
        bg_color = "F2F5F9" if idx % 2 == 1 else "FFFFFF"

        for c_idx, val in enumerate(vals):
            cell = table.cell(idx, c_idx)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            if c_idx in [0, 1]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.add_run(val)
            run.font.size = Pt(8.5)

    doc.save(docx_path)
    print(f"Generated Logbook (144 Days): {docx_path}")

# --- 3. GENERATE DAFTAR HADIR FILES ---
def generate_hadir_files(role, name, npm, hadir_fname):
    md_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.md"
    html_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.html"
    docx_path = f"c:\\Users\\pogoi\\Herd\\digiper\\{hadir_fname}.docx"

    # Markdown
    md_content = f"# DAFTAR HADIR MAGANG MAHASISWA\n"
    md_content += f"**Tahun Ajaran:** 2025/2026  \n"
    md_content += f"**Instansi:** Mahkamah Agung Republik Indonesia  \n"
    md_content += f"**Nama Mahasiswa:** {name}  \n"
    md_content += f"**NPM:** {npm}  \n"
    md_content += f"**Peran:** {role}  \n"
    md_content += f"**Total Kehadiran:** 144 Hari Hadir (100% Hadir / 1.152 Jam)  \n\n"
    md_content += "---\n\n"
    md_content += "| NO | NAMA MAHASISWA | NPM | TANGGAL HADIR | JUMLAH: H | JUMLAH: I | JUMLAH: S | JUMLAH: A | KETERANGAN |\n"
    md_content += "|:--:|:---|:---:|:---:|:--:|:--:|:--:|:--:|:---|\n"

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        md_content += f"| {idx} | {name} | {npm} | {day_name}, {date_str} | ✓ | | | | {work_mode} |\n"

    md_content += "\n---\n"
    md_content += "**Keterangan:**\n"
    md_content += "- **H**: Hadir (✓ Checkmark - 144 Hari)\n"
    md_content += "- **I**: Izin (0)\n"
    md_content += "- **S**: Sakit (0)\n"
    md_content += "- **A**: Alpa (0)\n\n"
    md_content += "**Catatan Kehadiran:**\n"
    md_content += "- **WFO (Work From Office)**: Hari Senin s.d. Kamis & Jumat Terjadwal (Kantor MA)\n"
    md_content += "- **WFH (Work From Home)**: Hari Sabtu & Jumat Terjadwal\n\n"
    md_content += f"Jakarta, {working_days[-1][1]}  \n"
    md_content += "**Instruktur / Pembimbing Magang Instansi:**  \n\n\n"
    md_content += "__________________________________________  \n"
    md_content += "NIP. …………………………………………  \n"

    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

    # HTML
    html_content = f"""<!DOCTYPE html>
<html lang="id">
<head>
    <meta charset="UTF-8">
    <title>DAFTAR HADIR MAGANG - {name}</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 30px; background-color: #f8f9fa; }}
        .container {{ max-width: 1100px; margin: 0 auto; background: #fff; padding: 30px; border-radius: 8px; box-shadow: 0 4px 12px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #1f497d; margin-bottom: 5px; font-size: 22px; text-decoration: underline; }}
        h2 {{ text-align: center; color: #555; margin-top: 0; font-size: 15px; font-weight: normal; margin-bottom: 25px; }}
        .info {{ margin-bottom: 20px; font-size: 14px; line-height: 1.6; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 13px; }}
        th {{ background-color: #1f497d; color: white; padding: 8px; text-align: center; border: 1px solid #1f497d; font-weight: bold; }}
        td {{ padding: 7px 8px; border: 1px solid #dce4ec; text-align: center; }}
        td.left {{ text-align: left; }}
        tr:nth-child(even) {{ background-color: #f2f5f9; }}
        .wfh {{ color: #d97706; font-weight: bold; }}
        .wfo {{ color: #15803d; font-weight: bold; }}
        .check {{ color: #15803d; font-weight: bold; font-size: 16px; }}
        .footer {{ margin-top: 30px; float: right; text-align: center; width: 280px; font-size: 14px; }}
        .clear {{ clear: both; }}
    </style>
</head>
<body>
<div class="container">
    <h1>DAFTAR HADIR MAGANG</h1>
    <h2>Tahun Ajaran 2025/2026 — Mahkamah Agung Republik Indonesia</h2>
    <div class="info">
        <b>Nama Mahasiswa:</b> {name}<br>
        <b>NPM:</b> {npm}<br>
        <b>Peran / Posisi:</b> {role}<br>
        <b>Total Kehadiran:</b> 144 Hari Hadir (100% Hadir / 1.152 Jam)
    </div>
    <table>
        <thead>
            <tr>
                <th rowspan="2" width="5%">NO</th>
                <th rowspan="2" width="25%">NAMA MAHASISWA</th>
                <th rowspan="2" width="13%">NPM</th>
                <th rowspan="2" width="22%">TANGGAL HADIR</th>
                <th colspan="4" width="16%">JUMLAH</th>
                <th rowspan="2" width="19%">KETERANGAN</th>
            </tr>
            <tr>
                <th width="4%">H</th>
                <th width="4%">I</th>
                <th width="4%">S</th>
                <th width="4%">A</th>
            </tr>
        </thead>
        <tbody>
"""
    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        cls = "wfh" if "WFH" in work_mode else "wfo"
        html_content += f"""            <tr>
                <td>{idx}</td>
                <td class="left">{name}</td>
                <td>{npm}</td>
                <td>{day_name}, {date_str}</td>
                <td class="check">✓</td>
                <td></td>
                <td></td>
                <td></td>
                <td class="{cls}">{work_mode}</td>
            </tr>\n"""

    html_content += f"""        </tbody>
    </table>
    <div style="margin-top: 20px; font-size: 13px; line-height: 1.6;">
        H : Hadir (✓ Checkmark - 144 Hari)<br>
        I : Izin (0)<br>
        S : Sakit (0)<br>
        A : Alpa (0)
    </div>
    <div class="footer">
        Jakarta, {working_days[-1][1]}<br>
        <b>Nama Instansi Magang: Mahkamah Agung RI</b><br><br><br><br>
        __________________________________________<br>
        <b>Instruktur Magang</b>
    </div>
    <div class="clear"></div>
</div>
</body>
</html>"""

    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    # DOCX
    doc = docx.Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DAFTAR HADIR MAGANG\n")
    r.bold = True
    r.underline = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    r2 = p.add_run("Tahun Ajaran 2025/2026 — Mahkamah Agung Republik Indonesia\n")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_after = Pt(8)
    p_meta.add_run(f"Nama Mahasiswa : {name}\nNPM             : {npm}\nPeran / Posisi  : {role}\nTotal Kehadiran : 144 Hari Hadir (100% Hadir / 1.152 Jam)")

    table = doc.add_table(rows=len(working_days) + 2, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="1F497D")

    col_widths = [Inches(0.4), Inches(1.8), Inches(1.0), Inches(1.4), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(1.6)]

    hdr_cells_r0 = table.rows[0].cells
    hdr_cells_r1 = table.rows[1].cells

    hdr_cells_r0[0].text = "NO"
    hdr_cells_r0[0].merge(hdr_cells_r1[0])
    
    hdr_cells_r0[1].text = "NAMA MAHASISWA"
    hdr_cells_r0[1].merge(hdr_cells_r1[1])

    hdr_cells_r0[2].text = "NPM"
    hdr_cells_r0[2].merge(hdr_cells_r1[2])

    hdr_cells_r0[3].text = "TANGGAL HADIR"
    hdr_cells_r0[3].merge(hdr_cells_r1[3])

    hdr_cells_r0[4].text = "JUMLAH"
    hdr_cells_r0[4].merge(hdr_cells_r0[5]).merge(hdr_cells_r0[6]).merge(hdr_cells_r0[7])

    hdr_cells_r1[4].text = "H"
    hdr_cells_r1[5].text = "I"
    hdr_cells_r1[6].text = "S"
    hdr_cells_r1[7].text = "A"

    hdr_cells_r0[8].text = "KETERANGAN"
    hdr_cells_r0[8].merge(hdr_cells_r1[8])

    for row in [table.rows[0], table.rows[1]]:
        for cell in row.cells:
            set_cell_background(cell, "1F497D")
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9)

    for idx, (day_name, date_str, work_mode, dt) in enumerate(working_days, 1):
        r_idx = idx + 1
        row_cells = table.rows[r_idx].cells
        bg_color = "F2F5F9" if idx % 2 == 1 else "FFFFFF"

        vals = [str(idx), name, npm, f"{day_name}, {date_str}", "✓", "", "", "", work_mode]

        for c_idx, val in enumerate(vals):
            cell = row_cells[c_idx]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=50, bottom=50, left=60, right=60)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = col_widths[c_idx]

            p = cell.paragraphs[0]
            if c_idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(val)
            run.font.size = Pt(8.5)

            if c_idx == 4:
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D)
            elif "WFH" in val:
                run.bold = True
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

    doc.add_paragraph()
    p_leg = doc.add_paragraph()
    p_leg.paragraph_format.space_after = Pt(15)
    r_leg = p_leg.add_run("H  : Hadir (✓ Checkmark - 144 Hari)\nI   : Izin\nS   : Sakit\nA  : Alpa")
    r_leg.font.size = Pt(9.5)

    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sig.paragraph_format.space_before = Pt(15)
    r_sig = p_sig.add_run(f"Jakarta, {working_days[-1][1]}\nNama instansi magang: Mahkamah Agung RI\n\n\n\n__________________________________________\nInstruktur Magang")
    r_sig.font.size = Pt(10)

    doc.save(docx_path)
    print(f"Generated Hadir DOCX (144 Days): {docx_path}")

if __name__ == '__main__':
    for role, name, npm, logbook_fname, hadir_fname in STUDENTS:
        generate_logbook_files(role, name, npm, logbook_fname)
        generate_hadir_files(role, name, npm, hadir_fname)
