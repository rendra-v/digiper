import re
import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = OxmlElement('w:tblBorders')
        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{b}')
            border.set(qn('w:val'), val)
            border.set(qn('w:sz'), sz)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            borders.append(border)
        tblPr[0].append(borders)

def markdown_to_docx(md_path, docx_path, title="LOGBOOK LAPORAN KEGIATAN MAGANG"):
    if not os.path.exists(md_path):
        print(f"File not found: {md_path}")
        return

    doc = docx.Document()

    # Set page orientation to Landscape for Logbook tables so all 8 columns fit nicely!
    section = doc.sections[0]
    if "LOGBOOK" in os.path.basename(md_path).upper():
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    else:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Set Normal Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    def process_table(lines_chunk):
        if not lines_chunk:
            return
        
        # Parse table markdown
        rows_data = []
        for l in lines_chunk:
            l_str = l.strip()
            if not l_str.startswith('|'):
                continue
            # Ignore divider row (e.g. |:---:|:---|)
            if re.match(r'^\|[\s:\-|\+]+\|$', l_str):
                continue
            cols = [c.strip() for c in l_str.split('|')[1:-1]]
            if cols:
                rows_data.append(cols)

        if not rows_data:
            return

        num_rows = len(rows_data)
        num_cols = max(len(r) for r in rows_data)

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="B0C4DE", sz="4")

        # Column widths for 8-column logbook
        col_widths = [Inches(0.4), Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.7), Inches(1.5), Inches(1.5), Inches(1.5)]

        for r_idx, row in enumerate(rows_data):
            for c_idx, val in enumerate(row):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    
                    if c_idx < len(col_widths):
                        cell.width = col_widths[c_idx]

                    # Header styling
                    if r_idx == 0:
                        set_cell_background(cell, "1F497D") # Dark Blue Header
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(val)
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(9.5)
                    else:
                        # Zebra striping
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F2F5F9")
                        else:
                            set_cell_background(cell, "FFFFFF")
                        
                        p = cell.paragraphs[0]
                        if c_idx == 0 or c_idx == 1:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        # Process inline bold/code formatting
                        clean_val = val.replace('**', '').replace('`', '')
                        run = p.add_run(clean_val)
                        run.font.size = Pt(9)

        doc.add_paragraph() # spacing

    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip('\n')
        line_str = line.strip()

        if line_str.startswith('|'):
            table_lines.append(line_str)
            in_table = True
            idx += 1
            continue
        else:
            if in_table:
                process_table(table_lines)
                table_lines = []
                in_table = False

        if not line_str:
            idx += 1
            continue

        # Headings
        if line_str.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_str[2:].replace('**', ''))
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line_str.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line_str[3:].replace('**', ''))
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line_str.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line_str[4:].replace('**', ''))
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif line_str.startswith('---'):
            idx += 1
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            clean_p = line_str.replace('**', '').replace('*', '')
            run = p.add_run(clean_p)
            run.font.size = Pt(10)

        idx += 1

    if in_table and table_lines:
        process_table(table_lines)

    doc.save(docx_path)
    print(f"Successfully generated: {docx_path}")

if __name__ == '__main__':
    base_dir = r"c:\Users\pogoi\Herd\digiper"
    files = [
        ("LOGBOOK_BE1_RIFA_REZA_FAHLEVI.md", "LOGBOOK_BE1_RIFA_REZA_FAHLEVI.docx"),
        ("LOGBOOK_BE2.md", "LOGBOOK_BE2.docx"),
        ("LOGBOOK_FE1.md", "LOGBOOK_FE1.docx"),
        ("LOGBOOK_FE2.md", "LOGBOOK_FE2.docx"),
        ("LAPORAN_MAGANG_DIGIPER.md", "LAPORAN_MAGANG_DIGIPER.docx")
    ]
    for md, docx_name in files:
        md_p = os.path.join(base_dir, md)
        docx_p = os.path.join(base_dir, docx_name)
        markdown_to_docx(md_p, docx_p)
